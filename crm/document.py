from docx import Document
from datetime import datetime
import os, io

def create_agreement_document(owner_details, agreement_details, notes, requestor_info):
    doc = Document()

    doc.add_heading('New Owner Agreement Preparation Request', level=1)

    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")

    doc.add_heading('Owner Details', level=2)
    doc.add_paragraph(f"Owner Name: {owner_details['name']}")
    doc.add_paragraph(f"Owner Type: {owner_details['type']}")
    doc.add_paragraph(f"Agreement Type: {owner_details['agreement_type']}")
    doc.add_paragraph(f"Contact Number: {owner_details['contact']}")
    doc.add_paragraph(f"Email: {owner_details['email']}")

    doc.add_heading('Agreement Details', level=2)

    if owner_details['agreement_type'] == 'Asset Management':
        doc.add_heading('For Asset Management Owners', level=3)
        doc.add_paragraph(f"Contract Mode: {agreement_details['contract_mode']}")
        if 'revenue_type' in agreement_details:
            doc.add_paragraph(f"Revenue Type: {agreement_details['revenue_type']}")
        doc.add_paragraph(f"Contract Status: {agreement_details['contract_status']}")
    
    elif owner_details['agreement_type'] == 'Brokerage':
        doc.add_heading('For Brokerage Owners', level=3)
        doc.add_paragraph(f"Contract Mode: {agreement_details['contract_mode']}")
        doc.add_paragraph(f"Contract Status: {agreement_details['contract_status']}")

    doc.add_heading('Notes/Instructions', level=2)
    doc.add_paragraph(notes)

    doc.add_heading('Request for Action', level=2)
    doc.add_paragraph(f"Please prepare the {owner_details['agreement_type']} agreement based on the provided details.")
    doc.add_paragraph(f"Kindly review and finalize the agreement for {owner_details['name']}.")
    doc.add_paragraph("Send the finalized agreement back to the admin upon completion.")

    doc.add_heading('Contact Information', level=2)
    doc.add_paragraph(f"Requestor Name: {requestor_info['name']}")
    doc.add_paragraph(f"Role: {requestor_info['role']}")
    doc.add_paragraph(f"Contact Information: {requestor_info['contact']}")

    filename = f"{owner_details['agreement_type'].replace(' ', '_')}_Agreement_Request.docx"
    file_path = os.path.join('agreements', filename)
    output = io.StringIO()
    print(f"{filename} created successfully!")
    document_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    output.close()

    return file_path, document_content
